from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm,Pt,RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.oxml.ns import qn
import time_month
#import category

def create_doc(year, month, res_list, depart_list):

    (month_list, days) = time_month.month_list(year, month)

    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # heading
    heading = document.add_heading()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_content = "疫苗接种点应急保障人员 （" + str(year) + "." + str(month) + ".1-" + str(month) + "." + str(days) + "） "
    h = heading.add_run(heading_content)
    h.font.color.rgb = RGBColor(0,0,0)
    h.font.name=u'宋体'
    h._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # table style
    cell_width = 5
    cell_height = 0.5

    # table add text method
    def add_text(table, row, col, text, font_size, if_bold):
        cell = table.cell(row, col)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(font_size)
        run.font.bold = if_bold
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # create table
    def create_table(document):
        table = document.add_table(rows=10,cols=8,style='Table Grid')
        document.add_paragraph("\n")

        for i in range(0,8):
            tcell = table.cell(0,i)
            tcell.width = Cm(cell_width) 
            shading = parse_xml(r'<w:shd {} w:fill="98FB98"/>'.format(nsdecls('w')))
            tcell._tc.get_or_add_tcPr().append(shading)
        table.rows[0].height = Cm(cell_height)

        # table add column text method
        def add_column_text(table, row_num, col, text_list):
            for i in range(1,row_num):
                add_text(table, i, col, text_list[i-1], 6.5, False)

        # table add place column
        add_column_text(table, 10, 0, ["稽东","福全","漓渚","王坛","平水","兰亭","方舱接种点","南大楼核酸采集","南大楼核酸采集"])

        return table

    def table_add_calendar(table, week_list, month):
        for col in range(1,8):
            if week_list[col-1] == 0:
                continue
            else:
                add_text(table, 0, col, str(month) + "." + str(week_list[col-1]), 8, True)
        return table

    def table_add_daily(table, week_d_list):
        for col in range(1,8):
            if len(week_d_list[col-1]) == 0:
                continue
            else:
                for row in range(1,7):
                    try:
                        add_text(table, row, col, depart_list[week_d_list[col-1][row-1]], 8, False)
                    except:
                        print(row, ",", col, "not valid")
        return table
                        

    tables = [create_table(document) for i in range(0,len(month_list))]
    tables = [table_add_calendar(tables[i], month_list[i], month) for i in range(len(month_list))]

    res_list = time_month.detail_month_list(year, month, res_list)
    
    tables = [table_add_daily(tables[i], res_list[i]) for i in range(len(res_list))]
    
    save_name = "绍兴二院新冠疫苗接种应急保障科室排班（"+str(year)+"."+str(month)+"）.docx"
    document.save(save_name)

